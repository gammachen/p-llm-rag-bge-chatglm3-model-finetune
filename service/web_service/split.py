import os

from pypdf import PdfReader
from docx import Document
import re
from tqdm import tqdm


def split_content_to_parse(content, max_length):
    """
    将文本内容按句子分割并重新组合成段落
    
    该函数会根据标点符号将文本分割成句子，然后将句子重新组合成不超过指定长度的段落。
    
    Args:
        content (str): 需要分割的原始文本内容
        max_length (int): 每个段落的最大字符长度
    
    Returns:
        tuple: 包含两个元素的元组
            - list: 分割后的段落列表
            - list: 每个段落中包含的句子数量列表
    """
    # 使用正则表达式按标点符号分割句子
    sentences = re.split(r"([。！？；.!?;])", content)
    sentences.append("")
    sentences = ["".join(i) for i in zip(sentences[0::2], sentences[1::2])]
    if sentences[-1] == "":
        sentences.pop(-1)
    
    # 初始化段落和句子计数器
    all_paras = []
    all_sentences_num_in_paras = []
    paras = []
    sentences_num_in_paras = 0
    sentences_num = len(sentences)
    
    # 遍历所有句子，将它们组合成适当长度的段落
    for idx, sen in enumerate(sentences):
        if len("".join(paras)) <= max_length:
            paras.append(sen)
            sentences_num_in_paras += 1
        if len("".join(paras)) > max_length:
            if sentences_num_in_paras > 1:
                # 如果当前段落包含多个句子，则移除最后一个句子并添加到结果中
                all_paras.append("".join(paras[:-1]))
                all_sentences_num_in_paras.append(sentences_num_in_paras - 1)
                paras = []
                sentences_num_in_paras = 1
                paras.append(sen)
            else:
                # 如果当前段落只有一个句子，则直接添加到结果中
                all_paras.append("".join(paras))
                all_sentences_num_in_paras.append(sentences_num_in_paras)
                paras = []
                sentences_num_in_paras = 0
        # 处理最后一个句子
        if idx == sentences_num - 1 and sentences_num_in_paras >= 1:
            all_paras.append("".join(paras))
            all_sentences_num_in_paras.append(sentences_num_in_paras)
    return all_paras, all_sentences_num_in_paras


def load_file(file_path, max_para_length=512):
    """
    加载并处理文件内容，支持PDF和DOCX格式
    
    该函数根据文件扩展名调用相应的处理函数，将文件内容分割成适当长度的段落
    
    Args:
        file_path (str): 文件路径
        max_para_length (int, optional): 每个段落的最大字符长度. Defaults to 512.
    
    Returns:
        list: 包含分割后段落的列表
    """
    max_para_length = max_para_length

    def _get_pdf_lines(pdf_path):
        """
        提取PDF文件中的文本内容并分割成段落
        
        Args:
            pdf_path (str): PDF文件路径
            
        Returns:
            list: 包含分割后段落的列表
        """
        reader = PdfReader(pdf_path)
        number_of_pages = len(reader.pages)
        all_paras = []
        print("Start loading pdf")
        result = []
        for i in tqdm(range(number_of_pages)):
            page = reader.pages[i]
            all_lines = page.extract_text()
            paras, _ = split_content_to_parse(all_lines, max_para_length)
            all_paras += paras
        return all_paras

    def _get_doc_lines(doc_path):
        """
        提取DOCX文件中的文本内容并分割成段落
        
        Args:
            doc_path (str): DOCX文件路径
            
        Returns:
            list: 包含分割后段落的列表
        """
        doc = Document(doc_path)
        all_paras = []
        print("Start loading doc")
        for paragraph in tqdm(doc.paragraphs):
            paras, _ = split_content_to_parse(paragraph.text, max_para_length)
            all_paras += paras
        return all_paras

    def load_files(file_path):
        """
        根据文件扩展名加载相应类型的文件
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            list: 包含分割后段落的列表
        """
        para = []
        if file_path.endswith("pdf"):
            para = _get_pdf_lines(file_path)
        elif file_path.endswith("docx"):
            para = _get_doc_lines(file_path)
        return para

    return load_files(file_path=file_path)


class LoadFileService:
    def __init__(self, max_para_length=512):
        self.max_para_length = max_para_length

    def load_files(self, file_path):
        para = []
        if file_path.endswith("pdf"):
            para = self._get_pdf_lines(file_path)
        elif file_path.endswith("docx"):
            para = self._get_doc_lines(file_path)
        return para

    def _get_pdf_lines(self, pdf_path):
        reader = PdfReader(pdf_path)
        number_of_pages = len(reader.pages)
        all_paras = []
        print("Start loading pdf")
        result = []
        for i in tqdm(range(number_of_pages)):
            page = reader.pages[i]
            all_lines = page.extract_text()
            paras, _ = split_content_to_parse(all_lines, self.max_para_length)
            all_paras += paras
        return all_paras

    def _get_doc_lines(self, doc_path):
        doc = Document(doc_path)
        all_paras = []
        print("Start loading doc")
        for paragraph in tqdm(doc.paragraphs):
            paras, _ = split_content_to_parse(paragraph.text, self.max_para_length)
            all_paras += paras
        return all_paras
